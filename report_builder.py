import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import List, Dict, Optional


def add_horizontal_line(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_heading_color(heading, r, g, b):
    for run in heading.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def add_section_heading(doc, title, level=1):
    h = doc.add_heading(title, level=level)
    if level == 1:
        set_heading_color(h, 31, 78, 121)
    elif level == 2:
        set_heading_color(h, 55, 138, 221)
    return h


def add_label_value(doc, label, value):
    p  = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(str(value))
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)


def set_cell_header(cell, text, bold=True, color=None):
    """Set table header text with bold/color without wiping runs."""
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def insert_image_safe(doc, image_path, caption=""):
    if not image_path or not os.path.exists(image_path):
        r = doc.add_paragraph().add_run("[Image Not Available]")
        r.italic = True
        r.font.color.rgb = RGBColor(150, 150, 150)
        return
    try:
        doc.add_picture(image_path, width=Inches(5.0))
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cp.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)
    except Exception as e:
        r = doc.add_paragraph().add_run(f"[Image could not be loaded: {e}]")
        r.italic = True
        r.font.color.rgb = RGBColor(150, 150, 150)


def get_images_for_area(area_name, image_refs, image_records, image_placement_map):
    id_to_path = {rec["image_id"]: rec["file_path"] for rec in image_records}
    matched = []

    for img_id in (image_refs or []):
        path = id_to_path.get(img_id)
        if path:
            matched.append({"id": img_id, "path": path, "desc": img_id})

    for entry in (image_placement_map or []):
        if entry.get("place_in_area", "").lower() == area_name.lower():
            img_id = entry.get("image_id", "")
            path   = id_to_path.get(img_id)
            if path and not any(m["id"] == img_id for m in matched):
                matched.append({"id": img_id, "path": path, "desc": entry.get("description", img_id)})

    return matched


def build_cover_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("DETAILED DIAGNOSTIC REPORT")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(31, 78, 121)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Property Inspection & Thermal Analysis")
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    add_horizontal_line(doc)
    doc.add_paragraph()

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(f"Report Generated: {datetime.now().strftime('%B %d, %Y')}")
    dr.font.size = Pt(11)
    dr.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()


def build_section_1(doc, summary):
    add_section_heading(doc, "1. Property Issue Summary")
    doc.add_paragraph(summary.get("overview", "Not Available"))
    doc.add_paragraph()

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers   = ["Total Issues", "Critical", "Moderate", "Minor"]
    values    = [str(summary.get(k, 0)) for k in ["total_issues_found","critical_count","moderate_count","minor_count"]]
    colors    = [(31,78,121),(192,0,0),(197,90,17),(31,120,55)]

    for i, (h, v, c) in enumerate(zip(headers, values, colors)):
        set_cell_header(table.rows[0].cells[i], h, bold=True, color=c)
        vc = table.rows[1].cells[i]
        vc.text = ""
        p  = vc.paragraphs[0]
        r  = p.add_run(v)
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(*c)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def build_section_2(doc, observations, image_records, image_placement_map):
    add_section_heading(doc, "2. Area-wise Observations")
    if not observations:
        doc.add_paragraph("Not Available")
        return

    for obs in observations:
        area = obs.get("area_name", "Unknown Area")
        add_section_heading(doc, area, level=2)
        add_label_value(doc, "Inspection findings", obs.get("inspection_findings", "Not Available"))
        add_label_value(doc, "Thermal findings",    obs.get("thermal_findings",    "Not Available"))

        p = doc.add_paragraph()
        p.add_run("Summary: ").bold = True
        p.add_run(obs.get("combined_summary", "Not Available"))

        imgs = get_images_for_area(area, obs.get("image_references", []), image_records, image_placement_map)
        if imgs:
            doc.add_paragraph()
            for img in imgs:
                insert_image_safe(doc, img["path"], caption=f"Figure: {img['desc']}")
        else:
            r = doc.add_paragraph().add_run("Image: Not Available")
            r.italic = True
            r.font.color.rgb = RGBColor(150, 150, 150)

        add_horizontal_line(doc)
        doc.add_paragraph()


def build_section_3(doc, root_causes):
    add_section_heading(doc, "3. Probable Root Cause")
    if not root_causes:
        doc.add_paragraph("Not Available")
        return
    for item in root_causes:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item.get("issue", "") + ": ")
        r.bold = True
        p.add_run(item.get("cause", "Not Available"))
        areas = item.get("affected_areas", [])
        if areas:
            ap = doc.add_paragraph()
            ap.paragraph_format.left_indent = Inches(0.4)
            ap.add_run("Affected areas: ").italic = True
            ap.add_run(", ".join(areas))
    doc.add_paragraph()


def build_section_4(doc, severity_list):
    add_section_heading(doc, "4. Severity Assessment")
    if not severity_list:
        doc.add_paragraph("Not Available")
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, col in enumerate(["Area", "Severity", "Reasoning"]):
        set_cell_header(table.rows[0].cells[i], col, bold=True)

    sev_colors = {"Critical": RGBColor(192,0,0), "Moderate": RGBColor(197,90,17), "Minor": RGBColor(31,120,55)}

    for item in severity_list:
        row = table.add_row().cells
        row[0].text = item.get("area_name", "")
        sev = item.get("severity", "Minor")
        row[1].text = ""
        sr = row[1].paragraphs[0].add_run(sev)
        sr.bold = True
        sr.font.color.rgb = sev_colors.get(sev, RGBColor(100,100,100))
        row[2].text = item.get("reasoning", "Not Available")

    doc.add_paragraph()


def build_section_5(doc, actions):
    add_section_heading(doc, "5. Recommended Actions")
    if not actions:
        doc.add_paragraph("Not Available")
        return

    order  = {"Immediate": 0, "Short-term": 1, "Long-term": 2}
    pcolors = {"Immediate": RGBColor(192,0,0), "Short-term": RGBColor(197,90,17), "Long-term": RGBColor(31,120,55)}
    sorted_actions = sorted(actions, key=lambda x: order.get(x.get("priority","Long-term"), 3))

    for action in sorted_actions:
        pri = action.get("priority", "Long-term")
        p   = doc.add_paragraph()
        r1  = p.add_run(f"[{pri}] ")
        r1.bold = True
        r1.font.color.rgb = pcolors.get(pri, RGBColor(100,100,100))
        r2 = p.add_run(f"{action.get('area','')} — ")
        r2.bold = True
        p.add_run(action.get("action", "Not Available"))

        urg = action.get("estimated_urgency", "")
        if urg:
            up = doc.add_paragraph()
            up.paragraph_format.left_indent = Inches(0.3)
            ur = up.add_run(f"Urgency: {urg}")
            ur.italic = True
            ur.font.size = Pt(10)
            ur.font.color.rgb = RGBColor(100,100,100)

    doc.add_paragraph()


def build_section_6(doc, notes):
    add_section_heading(doc, "6. Additional Notes")
    if not notes:
        doc.add_paragraph("No additional notes.")
        return
    for note in notes:
        doc.add_paragraph(style="List Bullet").add_run(str(note))
    doc.add_paragraph()


def build_section_7(doc, missing_info):
    add_section_heading(doc, "7. Missing or Unclear Information")
    if not missing_info:
        doc.add_paragraph("All expected information was found in the documents.")
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, col in enumerate(["Expected Field", "Status", "Detail"]):
        set_cell_header(table.rows[0].cells[i], col, bold=True)

    for item in missing_info:
        row = table.add_row().cells
        row[0].text = item.get("field", "")
        row[1].text = item.get("status", "Not Available")
        row[2].text = item.get("detail", "")

    doc.add_paragraph()


def build_ddr_report(ddr_data: dict, image_records: List[Dict], output_path: str) -> str:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    image_placement_map = ddr_data.get("image_placement_map") or []

    build_cover_page(doc)
    build_section_1(doc, ddr_data.get("property_issue_summary", {}))
    build_section_2(doc, ddr_data.get("area_wise_observations", []), image_records, image_placement_map)
    build_section_3(doc, ddr_data.get("probable_root_cause", []))
    build_section_4(doc, ddr_data.get("severity_assessment", []))
    build_section_5(doc, ddr_data.get("recommended_actions", []))
    build_section_6(doc, ddr_data.get("additional_notes", []))
    build_section_7(doc, ddr_data.get("missing_or_unclear_information", []))

    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc.save(output_path)
    print(f"Report saved: {output_path}")
    return output_path